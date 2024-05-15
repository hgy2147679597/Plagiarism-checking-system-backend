import logging
import os
import sys

import docx
import gensim
import jieba
from gensim.models.doc2vec import TaggedDocument


class Model:
    # 定义一个函数，将文档转换为 TaggedDocument 对象
    def __init__(self, docx_directory: str, model_file: str, test_docx_file: str, vector_size: int, min_count: int,
                 epochs: int):
        logging.basicConfig(level=logging.DEBUG)
        self.docx_directory = docx_directory
        self.model_file = model_file
        self.test_docx_file = test_docx_file
        self.vector_size = vector_size
        self.min_count = min_count
        self.epochs = epochs
        logging.info("Doc2Vec init")

    @staticmethod
    def process_docx_to_tagged_document(docx_file, doc_name):
        doc = docx.Document(docx_file)
        # 初始化一个空字符串，用于存储文本内容
        text_content = ""
        # 遍历文档中的每个段落并提取文本内容
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        symbols_to_remove = ['\n', ' ', '、']
        for symbol in symbols_to_remove:
            text_content = text_content.replace(symbol, '')

        # 使用结巴分词进行分词
        tokens = jieba.lcut(text_content)

        # 创建 TaggedDocument 对象
        tagged_doc = TaggedDocument(words=tokens, tags=[doc_name])
        return tagged_doc

    def train_or_load_model(self):
        if os.path.exists(self.model_file):
            # 如果模型文件存在，则加载模型
            model = gensim.models.doc2vec.Doc2Vec.load(self.model_file)
        else:
            # 存储 TaggedDocument 对象的列表
            train_corpus = self.get_train_corpus()
            # 创建和训练 Doc2Vec 模型
            model = gensim.models.doc2vec.Doc2Vec(vector_size=40, min_count=0, epochs=40)
            model.build_vocab(train_corpus)
            model.train(train_corpus, total_examples=model.corpus_count, epochs=model.epochs)
            logging.info('模型训练完毕')
            # 保存训练好的模型
            model.save(self.model_file)
        return model

    def start(self):
        try:
            model = self.train_or_load_model()
            if self.test_docx_file.endswith(".docx") or self.test_docx_file.endswith(".doc"):
                test_docx_path = os.path.join(self.docx_directory, self.test_docx_file)
                tagged_test_doc = self.process_docx_to_tagged_document(test_docx_path, self.test_docx_file)
                inferred_vector = model.infer_vector(tagged_test_doc.words)
                sims = model.dv.most_similar([inferred_vector], topn=len(model.dv))
                logging.info(f"与测试文档 {self.test_docx_file} 相似度：")
                logging.info('[Model.start]返回相似度查询结果: ', sims)
                return sims
        except Exception as e:
            raise e

    def train(self):
        try:
            # 存储 TaggedDocument 对象的列表
            train_corpus = self.get_train_corpus()

            # 创建和训练 Doc2Vec 模型
            model = gensim.models.doc2vec.Doc2Vec(vector_size=self.vector_size, min_count=self.min_count,
                                                  epochs=self.epochs)
            model.build_vocab(train_corpus)
            model.train(train_corpus, total_examples=model.corpus_count, epochs=model.epochs)
            print('[Model.train]模型训练完毕,将会保存在:', self.model_file)
            directory = os.path.dirname(self.model_file)
            # 创建目录（如果不存在）
            if not os.path.exists(directory):
                os.makedirs(directory)
            if not os.path.exists(self.model_file):
                print('[Model.train] 新建模型成功')
                with open(self.model_file, 'a') as f:
                    pass
            else:
                print('[Model.train] 新建模型失败,已存在模型')
            # 保存训练好的模型
            model.save(self.model_file)
        except Exception as e:
            # 捕获异常并抛出给调用者
            raise e

    def get_train_corpus(self):
        train_corpus = []
        # 遍历处理每个文档
        for doc_name, docx_file in enumerate(os.listdir(self.docx_directory)):
            if docx_file.startswith("m") and (docx_file.endswith(".docx") or docx_file.endswith(".doc")):
                docx_path = os.path.join(self.docx_directory, docx_file)
                tagged_doc = self.process_docx_to_tagged_document(docx_path, docx_file)
                train_corpus.append(tagged_doc)
        return train_corpus


if __name__ == '__main__':
    test=False
    if not test:
        sys.stdout.reconfigure(encoding='utf-8')
        current_directory = os.getcwd()
        logging.info('工作目录:' + current_directory)
        logging.error(sys.argv)
        if len(sys.argv) != 8:
            logging.error('使用参数数量' + str(len(sys.argv)) +
                          "  Usage: python my_python_script.py <function_name> <docx_directory> <model_file> <test_docx_file> "
                          "<vector_size> <min_count> <epochs>")
            sys.exit(1)

        function_name = sys.argv[1]
        docx_directory = sys.argv[2]
        model_file = sys.argv[3]
        test_docx_file = sys.argv[4]
        vector_size = int(sys.argv[5])
        min_count = int(sys.argv[6])
        epochs = int(sys.argv[7])
        model = Model(docx_directory, model_file, test_docx_file, vector_size, min_count, epochs)
        logging.info("Function Name:", function_name)
        logging.info("Docx Directory:", docx_directory)
        logging.info("Model File:", model_file)
        logging.info("Test Docx File:", test_docx_file)
        logging.info("Vector Size:", vector_size)
        logging.info("Min Count:", min_count)
        logging.info("Epochs:", epochs)
        if function_name == 'start':
            try:
                sims = model.start()
                print(sims)
            except Exception as e:
                logging.error('[model.start] fail:'+str(e))
                sys.exit(2)
        elif function_name == 'train':
            try:
                model.train()
            except Exception as e:
                logging.error('[model.train] fail:' + str(e))
                sys.exit(3)
        else:
            logging.info('Function Name is wrong: ' + function_name + 'only work for <start> and <train>')
            sys.exit(4)
    else:
        model = Model('../articles', '../modeltest/doc2vec_model.model', '../articles/1997～2002年中文体育核心期刊中的比较学校体育文献综述.docx', 40,
                   0, 40)
        model.train()
